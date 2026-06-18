from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Strategic Importance Premiums in Geopolitical Market Shocks - citation-edited.docx"
OUT = ROOT / "Strategic Importance Premiums in Geopolitical Market Shocks - APA7 citation clean.docx"
CITATION_LOG = ROOT / "citation_insertion_log.md"
REFERENCE_LOG = ROOT / "reference_list_audit_log.md"


APPEND_CITATIONS = {
    39: {
        "citation": " (Caldara & Iacoviello, 2022; Farrell & Newman, 2019)",
        "reason": "Supports the Chapter 1 opening claim that geopolitical competition, uncertainty, sanctions, trade restrictions, and coercive economic policies create market risk.",
    },
    40: {
        "citation": " (Blackwill & Harris, 2016; Mazzucato, 2013, 2021; Rodrik, 2004; Weiss, 2014)",
        "reason": "Supports claims about strategic industrial policy, state-backed support, subsidies, procurement, financing, and protection for strategic sectors.",
    },
    42: {
        "citation": " (Baldwin, 1985; Brown & Warner, 1985; MacKinlay, 1997; Rodrik, 2004)",
        "reason": "Supports the literature-mapping paragraph linking economic statecraft, industrial policy, and event-study market-reaction research.",
    },
    43: {
        "citation": " (Block & Keller, 2011; Fama, 1970; Fama et al., 1969; Mazzucato, 2013, 2021)",
        "reason": "Supports the core argument that state support can change expected cash flows, downside risk, and investor valuation.",
    },
    44: {
        "citation": " (Brown & Warner, 1985; Fama, 1970; MacKinlay, 1997)",
        "reason": "Supports the event-study design and use of short-window abnormal returns to evaluate policy-announcement reactions.",
    },
    45: {
        "citation": " (Gelman & Loken, 2014; Nosek et al., 2018; Simmons et al., 2011)",
        "reason": "Supports the pre-specified empirical strategy and avoidance of post hoc outcome interpretation.",
    },
    46: {
        "citation": " (Baldwin, 1985; Blackwill & Harris, 2016; Caldara & Iacoviello, 2022; Rodrik, 2004)",
        "reason": "Supports the dissertation contribution connecting geopolitical risk, geoeconomics, industrial policy, and financial-market interpretation.",
    },
    47: {
        "citation": " (Amsden, 2001; Chang, 2002; Mazzucato, 2013, 2021; Weiss, 2014)",
        "reason": "Supports scope-condition claims about strategic industries, security, technological leadership, and state support.",
    },
    55: {
        "citation": " (Caldara & Iacoviello, 2022; Fama, 1970)",
        "reason": "Supports the distinction between aggregate geopolitical risk measures and cross-sectional firm-level market reactions.",
    },
    56: {
        "citation": " (Baldwin, 1985; Mazzucato, 2013, 2021; Rodrik, 2004)",
        "reason": "Supports the added pathway from geopolitical competition to strategic importance, state support, and state-backed opportunity.",
    },
    61: {
        "citation": " (Baldwin, 1985; Farrell & Newman, 2019, 2023)",
        "reason": "Supports the claim that economic statecraft literature emphasizes state behaviour, coercive capacity, and target vulnerability.",
    },
    62: {
        "citation": " (Baldwin, 1985; Blackwill & Harris, 2016; Fama, 1970)",
        "reason": "Supports the bridge from economic statecraft to investor interpretation and financial-market reaction.",
    },
    68: {
        "citation": " (Brown & Warner, 1985; MacKinlay, 1997)",
        "reason": "Supports the distinction between long-run industrial-policy outcomes and immediate short-window financial-market reactions.",
    },
    69: {
        "citation": " (Rodrik, 2004; Fama, 1970)",
        "reason": "Supports the distinction between strategic importance as a scope condition and state support as the valuation-relevant mechanism.",
    },
    74: {
        "citation": " (Brown & Warner, 1985; Kothari & Warner, 2007; MacKinlay, 1997)",
        "reason": "Supports the claim that short-run investor pricing differs from long-run industrial-policy success or welfare evaluation.",
    },
    75: {
        "citation": " (Fama, 1970; Mazzucato, 2013, 2021; Rodrik, 2004)",
        "reason": "Supports treating credible state support as a market signal that may affect expected cash flows, capital costs, and downside risk.",
    },
    80: {
        "citation": " (Kothari & Warner, 2007; MacKinlay, 1997)",
        "reason": "Supports event-study limitations involving anticipation, event contamination, benchmark choice, and window selection.",
    },
    86: {
        "citation": " (Fama, 1970; Shiller, 2003)",
        "reason": "Supports the investor-interpretation step between state support and market reaction.",
    },
    88: {
        "citation": " (Brown & Warner, 1985; Shiller, 2003)",
        "reason": "Supports heterogeneous market reactions and the treatment of supportive, null, and theory-weakening outcomes as empirical possibilities.",
    },
}


TEXT_REPLACEMENTS = {
    64: ("Amsden (1989, 2001)", "Amsden (2001)", "Removed unapproved Amsden 1989 from narrative citation; retained approved Amsden 2001."),
    72: ("Amsden (1989, 2001)", "Amsden (2001)", "Removed unapproved Amsden 1989 from narrative citation; retained approved Amsden 2001."),
    78: ("Ball and Brown (1968) similarly show how accounting information is reflected in stock prices. ", "", "Removed unapproved Ball and Brown 1968 sentence because it is not in the approved bibliography."),
    90: ("Amsden, 1989, 2001", "Amsden, 2001", "Removed unapproved Amsden 1989 from parenthetical citation; retained approved Amsden 2001."),
    108: ("Amsden, 1989, 2001", "Amsden, 2001", "Removed unapproved Amsden 1989 from parenthetical citation; retained approved Amsden 2001."),
}


APPROVED_REFERENCES = [
    "Amsden, A. H. (2001). The rise of “the rest”: Challenges to the West from late-industrializing economies. Oxford University Press.",
    "Baldwin, D. A. (1985). Economic statecraft. Princeton University Press.",
    "Blackwill, R. D., & Harris, J. M. (2016). War by other means: Geoeconomics and statecraft. Belknap Press of Harvard University Press.",
    "Block, F., & Keller, M. R. (Eds.). (2011). State of innovation: The U.S. government's role in technology development. Paradigm Publishers.",
    "Brown, S. J., & Warner, J. B. (1985). Using daily stock returns: The case of event studies. Journal of Financial Economics, 14(1), 3–31.",
    "Caldara, D., & Iacoviello, M. (2022). Measuring geopolitical risk. American Economic Review, 112(4), 1194–1225.",
    "Campbell, J. Y., Lo, A. W., & MacKinlay, A. C. (1997). The econometrics of financial markets. Princeton University Press.",
    "Chang, H.-J. (2002). Kicking away the ladder: Development strategy in historical perspective. Anthem Press.",
    "Evans, P. (1995). Embedded autonomy: States and industrial transformation. Princeton University Press.",
    "Fama, E. F. (1970). Efficient capital markets: A review of theory and empirical work. Journal of Finance, 25(2), 383–417.",
    "Fama, E. F., Fisher, L., Jensen, M. C., & Roll, R. (1969). The adjustment of stock prices to new information. International Economic Review, 10(1), 1–21.",
    "Farrell, H., & Newman, A. L. (2019). Weaponized interdependence. International Security, 44(1), 42–79.",
    "Farrell, H., & Newman, A. L. (2023). Underground empire: How America weaponized the world economy. Henry Holt and Company.",
    "Gelman, A., & Loken, E. (2014). The statistical crisis in science. American Scientist, 102(6), 460–465.",
    "Gerschenkron, A. (1962). Economic backwardness in historical perspective. Harvard University Press.",
    "Hirschman, A. O. (1980). National power and the structure of foreign trade. University of California Press.",
    "Johnson, C. (1982). MITI and the Japanese miracle. Stanford University Press.",
    "Kothari, S. P., & Warner, J. B. (2007). Econometrics of event studies. In Handbook of corporate finance: Empirical corporate finance (Vol. 1, pp. 3–36). Elsevier.",
    "Luttwak, E. N. (1990). From geopolitics to geo-economics. The National Interest, 20, 17–23.",
    "MacKinlay, A. C. (1997). Event studies in economics and finance. Journal of Economic Literature, 35(1), 13–39.",
    "Mazzucato, M. (2013). The entrepreneurial state. Anthem Press.",
    "Mazzucato, M. (2021). Mission economy. Harper Business.",
    "Nosek, B. A., et al. (2018). The preregistration revolution. PNAS, 115(11), 2600–2606.",
    "Rodrik, D. (2004). Industrial policy for the twenty-first century. Harvard University.",
    "Shiller, R. J. (2003). From efficient markets theory to behavioural finance. Journal of Economic Perspectives, 17(1), 83–104.",
    "Simmons, J. P., Nelson, L. D., & Simonsohn, U. (2011). False-positive psychology. Psychological Science, 22(11), 1359–1366.",
    "Wade, R. (1990). Governing the market. Princeton University Press.",
    "Weiss, L. (2014). America Inc.? Cornell University Press.",
]


def clone_run_format(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    dst.style = src.style
    dst.font.name = src.font.name
    dst.font.size = src.font.size
    if src.font.color and src.font.color.rgb:
        dst.font.color.rgb = src.font.color.rgb


def append_citation(paragraph, citation):
    if citation in paragraph.text:
        return False
    last_text_run = None
    for candidate in reversed(paragraph.runs):
        if candidate.text:
            last_text_run = candidate
            break
    append_text = citation
    if last_text_run is not None and last_text_run.text.rstrip().endswith("."):
        trailing_spaces = last_text_run.text[len(last_text_run.text.rstrip()):]
        last_text_run.text = last_text_run.text.rstrip()[:-1] + trailing_spaces
        append_text = citation + "."
    run = paragraph.add_run(append_text)
    if last_text_run is not None:
        clone_run_format(last_text_run, run)
    return True


def replace_text_preserve_first_run(paragraph, old, new):
    if old not in paragraph.text:
        return False
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = paragraph.runs[0].text.replace(old, new)
        return True
    full = paragraph.text.replace(old, new)
    paragraph.runs[0].text = full
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def set_paragraph_text_preserve_style(paragraph, text):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    if text:
        paragraph.add_run(text)


def main():
    doc = Document(SOURCE)
    citation_entries = []
    replacement_entries = []

    for idx, item in APPEND_CITATIONS.items():
        inserted = append_citation(doc.paragraphs[idx], item["citation"])
        if inserted:
            citation_entries.append((idx, item["citation"].strip(), item["reason"]))

    for idx, (old, new, reason) in TEXT_REPLACEMENTS.items():
        changed = replace_text_preserve_first_run(doc.paragraphs[idx], old, new)
        if changed:
            replacement_entries.append((idx, old, new, reason))

    ref_start = 278
    old_nonblank = [p.text for p in doc.paragraphs[ref_start:] if p.text.strip()]
    for offset, ref in enumerate(APPROVED_REFERENCES):
        set_paragraph_text_preserve_style(doc.paragraphs[ref_start + offset], ref)
    for idx in range(ref_start + len(APPROVED_REFERENCES), len(doc.paragraphs)):
        set_paragraph_text_preserve_style(doc.paragraphs[idx], "")

    doc.save(OUT)

    citation_log = [
        "# Citation Insertion Log",
        "",
        f"Source DOCX: `{SOURCE.name}`",
        f"Revised DOCX: `{OUT.name}`",
        "",
        "## Inserted Citations",
    ]
    for idx, citation, reason in citation_entries:
        citation_log.append(f"- Paragraph {idx}: inserted `{citation}`. {reason}")
    citation_log.extend(["", "## Citation Standardization / Approved-List Cleanup"])
    for idx, old, new, reason in replacement_entries:
        citation_log.append(f"- Paragraph {idx}: `{old}` -> `{new}`. {reason}")
    CITATION_LOG.write_text("\n".join(citation_log) + "\n", encoding="utf-8")

    removed = [
        "Amsden, A. H. (1989) entry removed because it is not in the approved reference list.",
        "Ball, R., & Brown, P. (1968) entry removed because it is not in the approved reference list.",
        "Drezner, D. W. (1999) entry removed because it is not in the approved reference list.",
        "DOIs, subtitles, expanded edition notes, and publisher details were normalized to the exact approved reference text supplied by the user.",
    ]
    ref_log = [
        "# Reference-List Audit Log",
        "",
        f"Original nonblank reference-list paragraphs inspected: {len(old_nonblank)}",
        f"Approved APA7 references retained: {len(APPROVED_REFERENCES)}",
        "",
        "## Actions",
    ]
    ref_log.extend(f"- {entry}" for entry in removed)
    ref_log.extend(["", "## Final Approved Reference List"])
    ref_log.extend(f"- {ref}" for ref in APPROVED_REFERENCES)
    REFERENCE_LOG.write_text("\n".join(ref_log) + "\n", encoding="utf-8")

    print(OUT)
    print(CITATION_LOG)
    print(REFERENCE_LOG)


if __name__ == "__main__":
    main()
